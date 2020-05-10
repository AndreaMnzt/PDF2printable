from PIL import Image
import io
from docx import Document
try:
    from StringIO import StringIO as StringIO ## for Python 2
except ImportError:
    from io import BytesIO as StringIO  ## for Python 3
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from math import ceil
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#credit MadisonTrash https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def white_border(cell):
		set_cell_border(
			cell, top={"color": "#FFFFFF"}, bottom={"color": "#FFFFFF"},
			start={"color": "#FFFFFF"}, end={"color": "#FFFFFF"}
		)

def insert_images(document, sizes, imgs_per_page, right_page):
	n_rows, n_columns, side_bar, top_bar, img_width, img_height = sizes

	#get tables in docx document
	tables = document.tables

	#start by a 0x(n_columns+1) grid
	table = document.add_table(rows=1, cols=n_columns+1)
	table.rows[0].height = top_bar

	if right_page: #put the border in the left
		table.columns[0].width = side_bar
		for j in range(1,len(table.columns)):
			table.columns[j].width = int(ceil(img_width))

	else: #put the border in the right
		for j in range(len(table.columns)-1):
				table.columns[j].width = int(ceil(img_width))
		table.columns[-1].width = side_bar

	table.style = "Table Grid"
	table.autofit = False
	table.allow_autofit = False

	#if right_page:
	#	table.alignment = WD_TABLE_ALIGNMENT.LEFT
	#else:
	#	table.alignment = WD_TABLE_ALIGNMENT.LEFT

	#index for the column in the grid to insert the image
	col = 0

	#add a row to the table
	row_cells = table.add_row().cells
	row_cells[0].height = img_height

	#for every image
	for i in range(0,len(imgs_per_page)):

		#get the image to print in the table this cycle
		image = imgs_per_page[i]

		#every n_column image change row by creating a new row
		if(col == n_columns):
			col = 0
			row_cells = table.add_row().cells
			row_cells[0].height = img_height

		#change border accoring to right or left page
		if right_page:
			actual_col = col+1
		else:
			actual_col = col


		#create paragraph within the table
		paragraph = row_cells[actual_col].paragraphs[0]
		paragraph.style = 'List'
		cell_paragraph_format = paragraph.paragraph_format
		cell_paragraph_format.left_indent = Cm(0.1)
		cell_paragraph_format.right_indent = Cm(0)

		#set alignment of picture inside the table
		if right_page:
			paragraph.alignment = WD_TABLE_ALIGNMENT.LEFT
		else:
			paragraph.alignment = WD_TABLE_ALIGNMENT.RIGHT

		#add image to table
		run = paragraph.add_run()
		run.add_picture(StringIO(image), width = img_width , height = img_height)
		run.space_before = Cm(0)

		col += 1

	#set border to white
	for i in range(len(table.rows)):
		for j in range(len(table.columns)):
			white_border(table.cell(i,j))
